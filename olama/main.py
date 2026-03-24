from litellm import completion
from config import MODEL, API_BASE
import os
import docx

class User_chat_Ollama:
    def __init__(self):
        self.messages = []
        self.MODEL = MODEL
        self.API_BASE = API_BASE
        self.analze_files_in_directory()

    def add_question(self, question: str) -> str:
        self.messages.append({"role": "user", "content": question})
        response = completion(
            model=self.MODEL,
            messages=self.messages,
            api_base=self.API_BASE,
            request_timeout=360,
        )
        answer = response.choices[0].message.content
        self.messages.append({"role": "assistant", "content": answer})

        return answer

    def __str__(self):
        answer = ""
        for dicts in self.messages:
            if dicts["role"] == "assistant":
                answer += "Ответ от LLM: " + dicts['content'] + "\n"
            elif dicts["role"] == "user":
                answer += "Ваш вопрос: " + dicts['content'] + "\n"
        return answer

    def clear_chat_user(self, *args, **kwargs):
        if "number" in kwargs:
            self.messages.pop(kwargs['number'])
            self.messages.pop(kwargs['number'])
            return "Сообщение было удалено!"
        if "question" in kwargs:
            for dicts in self.messages:
                if dicts['content'] == kwargs['question']:
                    ind = self.messages.index(dicts)
                    self.messages.pop(ind)
                    self.messages.pop(ind)
                    return "Сообщение было удалено!"
        self.messages = []
        return "Весь чат был удален!"

    def read_docx_complete(self, file_path):
        try:
            doc = docx.Document(file_path)
            result = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    result.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = "\n".join([p.text for p in cell.paragraphs])
                        row_text.append(cell_text.strip())
                    if any(row_text):
                        result.append(" | ".join(row_text))

            return "\n".join(result)

        except Exception as e:
            return f"Ошибка: {e}"


    def analze_files_in_directory(self, *args):
        for i in os.listdir('data'):
            if ".txt" in i:
                file = open(f"data/{i}", encoding='utf-8')
                self.add_question(file.read())
            if ".docx" in i:
                self.add_question(self.read_docx_complete(f"data/{i}"))

        return True